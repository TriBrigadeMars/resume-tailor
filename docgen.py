"""Convert the AI-generated markdown-ish text into a clean .docx file."""

from __future__ import annotations

import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _apply_heading(doc, text: str, level: int):
    h = doc.add_heading(text, level=min(level, 3))
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def markdown_to_docx(markdown_text: str) -> Document:
    """Build a Word document from simple markdown (headings, bullets, text)."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        stripped = line.strip()
        # Headings: #, ##, ###
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            _apply_heading(doc, text, level)
        # Bullets
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(text)
        # Numbered lines: 1. Item, 1) Item, 10. Item
        else:
            numbered = re.match(r"^(\d+)[.)]\s+(.*)$", stripped)
            if numbered:
                text = numbered.group(2).strip()
                p = doc.add_paragraph(style="List Number")
                p.add_run(text)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run(stripped)

    return doc


def markdown_to_text(markdown_text: str) -> str:
    """Convert simple markdown to clean plain text (strip formatting markers).

    Shared by the Flask download route and the desktop native save path so both
    produce identical plain-text output.
    """
    lines = []
    for raw_line in markdown_text.splitlines():
        s = raw_line.strip()
        if not s:
            lines.append("")
            continue
        # Headings: strip the '#' markers.
        if s.startswith("#"):
            s = re.sub(r"^#+\s*", "", s)
        # Bullets: strip the '- ' / '* ' markers.
        elif s.startswith("- ") or s.startswith("* "):
            s = s[2:]
        # Numbered items: keep the number but normalize spacing after it.
        else:
            numbered = re.match(r"^(\d+)[.)]\s+(.*)$", s)
            if numbered:
                s = f"{numbered.group(1)}. {numbered.group(2).strip()}"
        lines.append(s)
    return "\n".join(lines).strip()
